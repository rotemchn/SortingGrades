import os
from docx import Document
from docx.shared import Inches

# This file makes a Word file from all the plots in this folder

# Function to get list of PNG files in directory
def get_png_files(directory):
    png_files = []
    for filename in os.listdir(directory):
        if filename.endswith(".png"):
            png_files.append(filename)
    return png_files


# Function to create Word document with images
def create_word_document(png_files):
    doc = Document()
    for png_file in png_files:
        # Add a page break before inserting each image
        if doc.paragraphs:
            doc.add_page_break()
        # Add title (image name without ".png")
        title = os.path.splitext(png_file)[0]
        doc.add_heading(title, level=1)
        # Insert image
        doc.add_picture(png_file, width=Inches(6.0))
    return doc


# Get list of PNG files
png_files = get_png_files("./")

# Create Word document
doc = create_word_document(png_files)

# Save the document
doc.save("images_document.docx")

